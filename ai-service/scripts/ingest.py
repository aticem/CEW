#!/usr/bin/env python3
"""
STRUCTURED DOCUMENT INGESTION SCRIPT for CEW AI Service

This script processes documents with structured parsing:
- Excel: Row-by-row key-value pairs with column headers
- PDF: Table extraction first, then text fallback
- Word: Meaningful paragraph grouping

Usage:
    cd ai-service
    python scripts/ingest.py

The script will:
1. Clear the existing ChromaDB collection
2. Scan the ./documents/ folder for supported files
3. Parse each document using structured extraction
4. Chunk the structured data
5. Generate embeddings using OpenAI
6. Store in ChromaDB with metadata
"""
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import time
import re
import uuid

# Fix Windows console encoding for Unicode
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        # If stdout cannot be reconfigured (rare), continue without crashing.
        pass

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.config import DOCUMENTS_DIR, CHUNK_SIZE, CHUNK_OVERLAP
from app.services.embedding_service import generate_embedding_sync
from app.services.chroma_service import add_documents, clear_collection, get_collection_stats
from app.utils.text_utils import chunk_text, clean_text

# Document parsers
import pdfplumber
import openpyxl
import pandas as pd
from docx import Document


def parse_excel_structured(filepath: Path) -> List[Dict[str, Any]]:
    """
    Parse Excel with structured row-by-row extraction.
    
    Each row becomes a structured key-value string:
    "SOURCE: filename | SHEET: sheet_name | DATA: col1: val1, col2: val2, ..."
    
    Args:
        filepath: Path to Excel file
        
    Returns:
        List of dicts with sheet, row_num, and structured text
    """
    results = []
    
    try:
        # Read all sheets with pandas
        excel_data = pd.read_excel(filepath, sheet_name=None, engine='openpyxl')
        
        for sheet_name, df in excel_data.items():
            # Skip empty sheets
            if df.empty:
                continue
            
            # Detect section_type from sheet name
            sheet_name_lower = sheet_name.lower()
            if "references" in sheet_name_lower or "reference" in sheet_name_lower:
                section_type = "references"
            else:
                section_type = "table"  # Excel sheets are typically tables
            
            # Get column names (headers)
            headers = df.columns.tolist()
            
            # Process each row
            for row_idx, row in df.iterrows():
                # Build structured key-value pairs
                data_parts = []
                
                for col_name in headers:
                    value = row[col_name]
                    
                    # Skip NaN/None values (empty cells)
                    if pd.isna(value):
                        continue
                    
                    # Convert to string and clean
                    value_str = str(value).strip()
                    
                    # Skip empty strings
                    if not value_str or value_str.lower() == 'nan':
                        continue
                    
                    # Add to data parts
                    data_parts.append(f"{col_name}: {value_str}")
                
                # Only create entry if we have data
                if data_parts:
                    structured_text = (
                        f"SOURCE: {filepath.name} | "
                        f"SHEET: {sheet_name} | "
                        f"ROW: {row_idx + 2} | "  # +2 because Excel rows start at 1 and header is row 1
                        f"DATA: {', '.join(data_parts)}"
                    )
                    
                    results.append({
                        "sheet": sheet_name,
                        "row_num": row_idx + 2,
                        "section_type": section_type,
                        "text": structured_text
                    })
        
        return results
        
    except Exception as e:
        print(f"      ⚠️  Excel parsing error: {e}")
        return []


def parse_pdf_structured(filepath: Path) -> List[Dict[str, Any]]:
    """
    Parse PDF with table extraction first, text fallback.
    
    Uses pdfplumber to:
    1. Try to extract tables (returns structured key-value format)
    2. Fall back to text extraction with layout preservation
    
    Args:
        filepath: Path to PDF file
        
    Returns:
        List of dicts with page, table_num (if any), and text
    """
    results = []
    pdf_kind = "UNKNOWN"
    text_page_count = 0
    total_pages = 0
    
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                total_pages += 1
                
                # Extract text for section detection
                page_text = page.extract_text() or ""
                page_text_lower = page_text.lower()
                
                # Detect section type based on page content
                section_type = None
                if page_num == 1:
                    section_type = "title"  # First page is title/intro
                elif "references" in page_text_lower or "reference" in page_text_lower:
                    # Check if this is actually a references section (not just mentioning the word)
                    if page_text_lower.count("references") > 2 or "reference" in page_text_lower[:200]:
                        section_type = "references"
                elif "table of contents" in page_text_lower or "contents" in page_text_lower[:100]:
                    section_type = "toc"
                
                # Try table extraction first
                tables = page.extract_tables()
                
                if tables:
                    # Process each table on the page
                    for table_idx, table in enumerate(tables, start=1):
                        if not table or len(table) < 2:  # Need at least header + 1 row
                            continue
                        
                        # First row is usually headers
                        headers = table[0]
                        
                        # Clean headers
                        headers = [
                            str(h).strip() if h else f"Column_{i}" 
                            for i, h in enumerate(headers)
                        ]
                        
                        # Process data rows
                        for row_idx, row in enumerate(table[1:], start=2):
                            data_parts = []
                            
                            for col_idx, (header, value) in enumerate(zip(headers, row)):
                                # Skip None/empty values
                                if value is None:
                                    continue
                                
                                value_str = str(value).strip()
                                if not value_str:
                                    continue
                                
                                data_parts.append(f"{header}: {value_str}")
                            
                            if data_parts:
                                structured_text = (
                                    f"SOURCE: {filepath.name} | "
                                    f"PAGE: {page_num} | "
                                    f"TABLE: {table_idx} | "
                                    f"ROW: {row_idx} | "
                                    f"DATA: {', '.join(data_parts)}"
                                )
                                
                                # Table chunks get section_type="table" (unless already set to references/toc)
                                chunk_section_type = section_type if section_type in ("references", "toc") else "table"
                                
                                results.append({
                                    "page": page_num,
                                    "table_num": table_idx,
                                    "section_type": chunk_section_type,
                                    "text": structured_text
                                })
                
                # If no tables found or tables didn't cover the whole page, extract text
                if page_text:
                    if len(page_text.strip()) > 200:
                        text_page_count += 1
                    # Clean and structure the text
                    text = clean_text(page_text)
                    
                    # Only add if not empty and we haven't already added tables
                    if text.strip() and (not tables or len(text) > 200):
                        # Add source metadata
                        structured_text = (
                            f"SOURCE: {filepath.name} | "
                            f"PAGE: {page_num} | "
                            f"CONTENT:\n{text}"
                        )
                        
                        results.append({
                            "page": page_num,
                            "section_type": section_type,  # Can be None, "title", "references", or "toc"
                            "text": structured_text
                        })
        
        # Classify PDF kind based on text density
        if total_pages > 0:
            text_ratio = text_page_count / total_pages
            if text_ratio > 0.8:
                pdf_kind = "TEXT_PDF"
            elif text_ratio < 0.2:
                pdf_kind = "SCANNED_PDF"
            else:
                pdf_kind = "DRAWING_PDF"

        # Attach pdf_kind to each result for downstream metadata
        for r in results:
            r["pdf_kind"] = pdf_kind
        
        return results
        
    except Exception as e:
        print(f"      ⚠️  PDF parsing error: {e}")
        return []


def parse_word_structured(filepath: Path) -> List[Dict[str, Any]]:
    """
    Parse Word document with meaningful paragraph grouping.
    
    Groups:
    - Headers (detected by style or formatting)
    - Paragraphs under each header
    - Tables (extracted as structured data)
    
    Args:
        filepath: Path to Word file
        
    Returns:
        List of dicts with section info and text
    """
    results = []
    
    try:
        doc = Document(filepath)
        
        current_section = "Introduction"
        section_paragraphs = []
        section_num = 0
        is_first_section = True
        
        def _detect_section_type(section_name: str, section_num: int) -> str | None:
            """Detect section_type from section name."""
            section_lower = section_name.lower()
            if section_num == 0 or is_first_section:
                return "title"
            elif "references" in section_lower or "reference" in section_lower:
                return "references"
            elif "table of contents" in section_lower or "contents" in section_lower:
                return "toc"
            return None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect headers (heuristics)
            is_header = False
            
            # Check if it's a heading style
            if para.style.name.startswith('Heading'):
                is_header = True
            # Check for all caps short lines
            elif text == text.upper() and len(text) < 100 and len(text) > 3:
                is_header = True
            # Check for numbered sections
            elif re.match(r'^\d+(\.\d+)*\s+[A-Z]', text):
                is_header = True
            
            if is_header:
                # Save previous section if it has content
                if section_paragraphs:
                    section_text = "\n\n".join(section_paragraphs)
                    structured_text = (
                        f"SOURCE: {filepath.name} | "
                        f"SECTION: {current_section} | "
                        f"CONTENT:\n{section_text}"
                    )
                    
                    section_type = _detect_section_type(current_section, section_num)
                    
                    results.append({
                        "section": current_section,
                        "section_num": section_num,
                        "section_type": section_type,
                        "text": structured_text
                    })
                    
                    section_num += 1
                    is_first_section = False
                
                # Start new section
                current_section = text
                section_paragraphs = []
            else:
                # Add to current section
                section_paragraphs.append(text)
        
        # Save final section
        if section_paragraphs:
            section_text = "\n\n".join(section_paragraphs)
            structured_text = (
                f"SOURCE: {filepath.name} | "
                f"SECTION: {current_section} | "
                f"CONTENT:\n{section_text}"
            )
            
            section_type = _detect_section_type(current_section, section_num)
            
            results.append({
                "section": current_section,
                "section_num": section_num,
                "section_type": section_type,
                "text": structured_text
            })
        
        # Parse tables
        for table_idx, table in enumerate(doc.tables, start=1):
            rows = []
            
            # Get headers from first row
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Process data rows
            for row_idx, row in enumerate(table.rows[1:], start=2):
                data_parts = []
                
                for header, cell in zip(headers, row.cells):
                    value = cell.text.strip()
                    if value:
                        data_parts.append(f"{header}: {value}")
                
                if data_parts:
                    structured_text = (
                        f"SOURCE: {filepath.name} | "
                        f"TABLE: {table_idx} | "
                        f"ROW: {row_idx} | "
                        f"DATA: {', '.join(data_parts)}"
                    )
                    
                    results.append({
                        "table_num": table_idx,
                        "row_num": row_idx,
                        "section_type": "table",
                        "text": structured_text
                    })
        
        return results
        
    except Exception as e:
        print(f"      ⚠️  Word parsing error: {e}")
        return []


def process_document(filepath: Path) -> List[Dict[str, Any]]:
    """
    Process a document with structured parsing.
    
    Args:
        filepath: Path to document
        
    Returns:
        List of chunk dicts ready for ChromaDB
    """
    suffix = filepath.suffix.lower()
    doc_name = filepath.name
    doc_type = "unknown"
    pdf_kind = None
    
    print(f"      📄 Parsing: {doc_name}")
    
    # Parse based on file type with structured extraction
    try:
        if suffix == ".pdf":
            doc_type = "pdf"
            sections = parse_pdf_structured(filepath)
            print(f"         ✓ Extracted {len(sections)} structured sections (tables + text)")
            if sections:
                pdf_kind = sections[0].get("pdf_kind")
        elif suffix in {".xlsx", ".xls"}:
            doc_type = "excel"
            sections = parse_excel_structured(filepath)
            print(f"         ✓ Extracted {len(sections)} rows as structured data")
        elif suffix == ".docx":
            doc_type = "word"
            sections = parse_word_structured(filepath)
            print(f"         ✓ Extracted {len(sections)} sections (paragraphs + tables)")
        else:
            print(f"         ⚠️  Unsupported file type: {suffix}")
            return []
    except Exception as e:
        print(f"         ✗ Parse error: {e}")
        return []
    
    if not sections:
        print(f"         ⚠️  No content extracted")
        return []
    
    # Chunk each section
    chunks = []
    for section in sections:
        text = section["text"]
        
        # Don't chunk structured data that's already concise
        # (tables/rows should stay together)
        if len(text.split()) < CHUNK_SIZE:
            # Keep as single chunk
            text_chunks = [text]
        else:
            # Chunk longer content
            text_chunks = chunk_text(text, CHUNK_SIZE, CHUNK_OVERLAP)
        
        for i, chunk_text_content in enumerate(text_chunks):
            chunk_data = {
                "doc_name": doc_name,
                "chunk_index": i,
                "text": chunk_text_content
            }
            
            # Add doc_type/pdf_kind metadata for downstream filtering/analytics
            if doc_type:
                chunk_data["doc_type"] = doc_type
            if pdf_kind:
                chunk_data["pdf_kind"] = pdf_kind
            
            # Add metadata from section
            if "page" in section:
                chunk_data["page"] = section["page"]
            if "sheet" in section:
                chunk_data["sheet"] = section["sheet"]
            if "section" in section:
                chunk_data["section"] = section["section"]
            if "table_num" in section:
                chunk_data["table_num"] = section["table_num"]
            if "row_num" in section:
                chunk_data["row_num"] = section["row_num"]
            # Add section_type metadata (for intent-aware boosting)
            if "section_type" in section:
                chunk_data["section_type"] = section["section_type"]
            
            chunks.append(chunk_data)
    
    print(f"         ✓ Created {len(chunks)} chunks")
    return chunks


def ingest_documents():
    """Main ingestion function with structured parsing."""
    print("=" * 70)
    print("  CEW AI-SERVICE — STRUCTURED DOCUMENT INGESTION")
    print("=" * 70)
    
    # Step 1: Clear existing collection (CRITICAL - prevents duplicate ID errors)
    print("\n[1/4] Clearing existing ChromaDB collection...")
    try:
        clear_collection()
        print("      ✓ Collection cleared and reset")
    except Exception as e:
        print(f"      ⚠️  Error clearing collection: {e}")
        print(f"      Continuing anyway, but duplicate IDs may occur...")
    
    # Step 2: Find all documents
    print(f"\n[2/4] Scanning {DOCUMENTS_DIR} for documents...")
    
    supported_extensions = {".pdf", ".xlsx", ".xls", ".docx"}
    
    if not DOCUMENTS_DIR.exists():
        DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
        print(f"      ⚠️  Created empty documents folder: {DOCUMENTS_DIR}")
        print(f"      ⚠️  Place your documents there and re-run this script.")
        return
    
    files = [
        f for f in DOCUMENTS_DIR.iterdir() 
        if f.is_file() and f.suffix.lower() in supported_extensions
    ]
    
    if not files:
        print(f"      ⚠️  No supported documents found in {DOCUMENTS_DIR}")
        print(f"      Supported formats: PDF, XLSX, XLS, DOCX")
        print(f"      Place your documents there and re-run this script.")
        return
    
    print(f"      Found {len(files)} document(s):")
    for f in files:
        size_kb = f.stat().st_size / 1024
        print(f"      - {f.name} ({size_kb:.1f} KB)")
    
    # Step 3: Process each document with structured parsing
    print(f"\n[3/4] Processing documents with structured parsing...")
    all_chunks = []
    
    for file in files:
        print(f"\n   📁 {file.name}")
        start_time = time.time()
        
        try:
            chunks = process_document(file)
            all_chunks.extend(chunks)
            
            elapsed = time.time() - start_time
            print(f"      ✅ Complete ({elapsed:.1f}s)")
            
        except Exception as e:
            print(f"      ✗ Error: {str(e)}")
    
    if not all_chunks:
        print("\n      ⚠️  No text content extracted from documents.")
        print("      This may happen if PDFs are scanned images (OCR not supported).")
        return
    
    print(f"\n      📊 Total chunks to embed: {len(all_chunks)}")
    
    # Step 4: Generate embeddings and store
    print(f"\n[4/4] Generating embeddings and storing in ChromaDB...")
    
    documents = []
    total_chunks = len(all_chunks)
    
    for i, chunk in enumerate(all_chunks):
        # Progress indicator
        if (i + 1) % 10 == 0 or i == 0 or i == total_chunks - 1:
            progress = ((i + 1) / total_chunks) * 100
            print(f"      Processing chunk {i + 1}/{total_chunks} ({progress:.0f}%)...")
        
        # Generate embedding
        try:
            embedding = generate_embedding_sync(chunk["text"])
        except Exception as e:
            print(f"      ⚠️  Embedding error for chunk {i}: {str(e)}")
            continue
        
        # Create unique ID with UUID suffix to prevent duplicates
        page_or_sheet = (
            chunk.get("page") or 
            chunk.get("sheet") or 
            chunk.get("section") or 
            "0"
        )
        # Add random UUID suffix to guarantee uniqueness
        unique_suffix = uuid.uuid4().hex[:8]
        doc_id = f"{chunk['doc_name']}_{page_or_sheet}_{chunk['chunk_index']}_{unique_suffix}"
        
        # Build metadata
        metadata = {
            "doc_name": chunk["doc_name"],
            "chunk_index": chunk["chunk_index"]
        }
        
        # Add optional metadata
        for key in ["page", "sheet", "section", "table_num", "row_num", "doc_type", "pdf_kind"]:
            if key in chunk:
                metadata[key] = chunk[key]
        
        documents.append({
            "id": doc_id,
            "text": chunk["text"],
            "embedding": embedding,
            "metadata": metadata
        })
    
    # Batch add to ChromaDB with duplicate ID check
    if documents:
        # Pre-submission validation: check for duplicate IDs
        doc_ids = [doc["id"] for doc in documents]
        unique_ids = set(doc_ids)
        
        if len(doc_ids) != len(unique_ids):
            print(f"      ⚠️  WARNING: Found {len(doc_ids) - len(unique_ids)} duplicate IDs!")
            print(f"      This should not happen with UUID suffixes.")
            print(f"      Proceeding anyway (ChromaDB will reject duplicates)...")
        else:
            print(f"      ✓ All {len(doc_ids)} IDs are unique")
        
        try:
            add_documents(documents)
            print(f"      ✓ Successfully stored {len(documents)} chunks in ChromaDB")
        except Exception as e:
            print(f"      ✗ ChromaDB error: {e}")
            print(f"      💡 Tip: Try deleting chroma_db/ folder and re-run")
            return
    
    # Final stats
    try:
        stats = get_collection_stats()
        
        print(f"\n" + "=" * 70)
        print(f"  ✅ STRUCTURED INGESTION COMPLETE")
        print(f"=" * 70)
        print(f"  Documents processed:  {len(files)}")
        print(f"  Chunks created:       {len(documents)}")
        print(f"  ChromaDB documents:   {stats.get('count', 0)}")
        print(f"  Parsing mode:         STRUCTURED (key-value pairs)")
        print(f"=" * 70)
    except Exception as e:
        print(f"\n  ⚠️  Could not retrieve final stats: {e}")


if __name__ == "__main__":
    ingest_documents()
